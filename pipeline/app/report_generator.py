from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime, timezone

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)


def create_disk_report(disk_data, patient_info, seg_sagittal=None, seg_axial=None, output_filename=None):
    """
    Generates a DOCX report with split tables and flipped images.
    """
    if not output_filename:
        output_filename = "mri_disk_report.docx"

    output_path = os.path.join(REPORTS_DIR, output_filename)

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Заключение по результатам МРТ поясничного отдела позвоночника')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    font_element = style._element.rPr
    font_element.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Add patient info block
    patient_lines = [
        f"Пациент: {patient_info.get('patient_name', 'N/A')}",
        f"Возраст: {patient_info.get('patient_age', 'N/A')}",
        f"Дата исследования: {patient_info.get('study_date', 'N/A')}",
        f"Методика: {patient_info.get('modality', 'MRI')}",
        f"Уровни сканирования: {patient_info.get('series_description', 'Lumbar Spine')}",
        f"Дата формирования отчета: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ]

    for line in patient_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # spacing

    if not disk_data:
        doc.add_paragraph("No data available.")
        doc.save(output_path)
        return output_path

    # === Table 1: Structural Analysis ===
    headers1 = [
        'Диск', 'Modic', 'Верхняя замыкательная пластина', 'Нижняя замыкательная пластина',
        'Спондилолизтез', 'Грыжа диска', 'Сужение диска', 'Протрузия диска', 'Степень Pfirrmann'
    ]
    keys1 = [
        'disk_label', 'Modic', 'UP endplate', 'LOW endplate', 'Spondylolisthesis',
        'Disc herniation', 'Disc narrowing', 'Disc bulging', 'Pfirrmann grade'
    ]

    # === Table 2: Detailed Metrics ===
    headers2 = [
        'Диск', 'Грыжа', 'Объем грыжи (mm³)', 'Hernia Max Protrusion (mm)',
        'Spondy Detected', 'Spondy Displacement (mm)', 'Spondy Displacement (%)',
        'Spondy Grade'
    ]
    keys2 = [
        'disk_label', 'hernia_detected', 'hernia_volume_mm3', 'hernia_max_protrusion_mm',
        'spondy_detected', 'spondy_displacement_mm', 'spondy_displacement_percentage',
        'spondy_grade'
    ]

    disk_mapping = {
        63: 'C2-C3', 64: 'C3-C4', 65: 'C4-C5', 66: 'C5-C6', 67: 'C6-C7',
        71: 'C7-T1', 72: 'T1-T2', 73: 'T2-T3', 74: 'T3-T4', 75: 'T4-T5',
        76: 'T5-T6', 77: 'T6-T7', 78: 'T7-T8', 79: 'T8-T9', 80: 'T9-T10',
        81: 'T10-T11', 82: 'T11-T12', 91: 'T12-L1', 92: 'L1-L2', 93: 'L2-L3',
        94: 'L3-L4', 95: 'L4-L5', 96: 'L5-S1', 100: 'L5-S1'
    }

    def format_cell_value(value, key=None):
        categorical_keys = {
            'UP endplate',
            'LOW endplate',
            'Spondylolisthesis',
            'Disc herniation',
            'Disc narrowing',
            'Disc bulging'
        }

        if value is None:
            return ""
        elif key in categorical_keys:
            if value == 1 or value is True:
                return "Есть"
            elif value == 0 or value is False:
                return "Нет"
            else:
                return str(value)  # fallback if unexpected value
        elif isinstance(value, bool):
            return "нет" if value is False else str(value)  # legacy bool handling
        elif key == 'Pfirrmann grade' and isinstance(value, (int, float)):
            return str(int(value) + 1)  # Increase by 1
        elif isinstance(value, list):
            return str(value).replace("[", "").replace("]", "")
        else:
            return str(value)

    # === Table 1 ===
    doc.add_paragraph("1. Структурный анализ по сегментам")
    table1 = doc.add_table(rows=1, cols=len(headers1))
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    for i, header in enumerate(headers1):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True

    for item in disk_data:
        row_cells = table1.add_row().cells
        for i, key in enumerate(keys1):
            raw_value = item.get(key, None)
            if key == 'disk_label' and raw_value in disk_mapping:
                value = disk_mapping[raw_value]
            else:
                value = raw_value
            text = format_cell_value(value, key)
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # === Table 2 ===
    doc.add_paragraph("2. Детальные данные обследования")
    table2 = doc.add_table(rows=1, cols=len(headers2))
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, header in enumerate(headers2):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True

    for item in disk_data:
        row_cells = table2.add_row().cells
        for i, key in enumerate(keys2):
            raw_value = item.get(key, None)
            if key == 'disk_label' and raw_value in disk_mapping:
                value = disk_mapping[raw_value]
            else:
                value = raw_value
            text = format_cell_value(value, key)
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # === Add Sagittal Segmentation Image (Рис. 27) ===
    if seg_sagittal is not None:
        if seg_sagittal.ndim == 3 and seg_sagittal.shape[0] == 3:
            seg_sagittal = np.transpose(seg_sagittal, (1, 2, 0))
        # Flip 180 degrees
        seg_sagittal = np.rot90(seg_sagittal, k=5)

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            plt.figure(figsize=(6, 4))
            plt.imshow(seg_sagittal)
            plt.axis('off')
            plt.title("Срез 27", fontsize=12, fontname='Times New Roman', fontweight='bold')
            plt.tight_layout()
            plt.savefig(tmpfile.name, dpi=150, bbox_inches='tight', pad_inches=0)
            plt.close()

            doc.add_paragraph()
            doc.add_picture(tmpfile.name, width=Inches(6))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Срез 27")
            caption_run.font.name = 'Times New Roman'
            caption_run.font.size = Pt(12)
            os.unlink(tmpfile.name)

    # === Add Axial Segmentation Image (Рис. 28) ===
    if seg_axial is not None:
        if seg_axial.ndim == 3 and seg_axial.shape[0] == 3:
            seg_axial = np.transpose(seg_axial, (1, 2, 0))
        # Flip 180 degrees
        seg_axial = np.rot90(seg_axial, k=5)

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            plt.figure(figsize=(6, 4))
            plt.imshow(seg_axial)
            plt.axis('off')
            plt.title("Срез 28", fontsize=12, fontname='Times New Roman', fontweight='bold')
            plt.tight_layout()
            plt.savefig(tmpfile.name, dpi=150, bbox_inches='tight', pad_inches=0)
            plt.close()

            doc.add_paragraph()
            doc.add_picture(tmpfile.name, width=Inches(6))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Срез 28")
            caption_run.font.name = 'Times New Roman'
            caption_run.font.size = Pt(12)
            os.unlink(tmpfile.name)

    doc.save(output_path)
    return output_path